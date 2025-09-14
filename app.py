import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from fpdf import FPDF
from io import BytesIO

# --- Page Config ---
st.set_page_config(page_title="Class Attendance", layout="centered")

# --- Student List ---
student_list = [
    { "Name": "BAIG ASFIYA MUSHARRAF", "Roll Number": "201" },
    { "Name": "BATTINWAR SHRUTI RAJAN", "Roll Number": "202" },
    { "Name": "BHISE RANJANA ROHIDAS", "Roll Number": "203" },
    { "Name": "BOKDE DIVYA GANESH", "Roll Number": "204" },
    { "Name": "CHAUDHARI SONALI HEMANTRAO", "Roll Number": "205" },
    { "Name": "DOLAI KHUSHI CHANDAN", "Roll Number": "206" },
    { "Name": "GABHANE MRUDULA PRAMOD", "Roll Number": "207" },
    { "Name": "GUPTA ACHAL SURESH", "Roll Number": "208" },
    { "Name": "GUPTA KRUPA SHANKAR", "Roll Number": "209" },
    { "Name": "KHAROLE ASMITA TILAK", "Roll Number": "210" },
    { "Name": "KULHAR KRUTIKA HIRALAL", "Roll Number": "211" },
    { "Name": "MOHADIKAR GUNJAN NARAYAN", "Roll Number": "212" },
    { "Name": "MOTHARKAR TRIVENI ISHWAR", "Roll Number": "213" },
    { "Name": "PIMPALKHEDE MAITHILY PRADEEP", "Roll Number": "214" },
    { "Name": "PRAJAPATI PRIYANKA JAGJEET", "Roll Number": "215" },
    { "Name": "RAKSHAK SONAL RAJKUMAR", "Roll Number": "216" },
    { "Name": "RAUT PURVA JOGENDER", "Roll Number": "217" },
    { "Name": "YADAV SHIVANGI SHIVANGI", "Roll Number": "218" },
    { "Name": "SINGH TANNU DHANANJAY", "Roll Number": "219" },
    { "Name": "HARIJAN SONI VIJAYKUMAR", "Roll Number": "220" },
    { "Name": "YELE PAYAL MADHU", "Roll Number": "221" },
    { "Name": "BAGHEL SAHIL SINGH", "Roll Number": "222" },
    { "Name": "BATBARVE DHIRAJ DILIP", "Roll Number": "223" },
    { "Name": "BHOYAR PRAFUL RAJU", "Roll Number": "224" },
    { "Name": "BRAMHANE SUGAT JAYANT", "Roll Number": "225" },
    { "Name": "CHAVAN HARIHAR RAVINDRA", "Roll Number": "226" },
    { "Name": "CHOUDHARY HUMENDRA KHEMRAJ", "Roll Number": "227" },
    { "Name": "DAMBHARE ROHIT SANTOSH", "Roll Number": "228" },
    { "Name": "DESHMANE GANESH SHRAVAN", "Roll Number": "229" },
    { "Name": "DHARMALE ANSHUL VIJAY", "Roll Number": "230" },
    { "Name": "GAYDHANE SHREYANSH GOPAL", "Roll Number": "231" },
    { "Name": "INGALE NISHANT NARENDRARAO", "Roll Number": "232" },
    { "Name": "JAISWAL SLOK SHRIKANT", "Roll Number": "233" },
    { "Name": "KASHIKAR PARTH HEMANT", "Roll Number": "234" },
    { "Name": "KAWALE VEDANT RAJESH", "Roll Number": "235" },
    { "Name": "KHAPRE AKSHAY ASHOK", "Roll Number": "236" },
    { "Name": "KONER KUSHOL ASIT", "Roll Number": "237" },
    { "Name": "KURVE MOHIT KRUSHNAKUMAR", "Roll Number": "238" },
    { "Name": "MESHRAM KRUNAL SANJAY", "Roll Number": "239" },
    { "Name": "MOON VANSH VILAS", "Roll Number": "240" },
    { "Name": "NAGPURE SAHIL SURESH", "Roll Number": "241" },
    { "Name": "PUNDE AMIT YADORAO", "Roll Number": "242" },
    { "Name": "RAMPURE NISHAD VILAS", "Roll Number": "243" },
    { "Name": "RATHOD GAURAV RAMESHWAR", "Roll Number": "244" },
    { "Name": "SELOKAR RUPESH RAJESH", "Roll Number": "245" },
    { "Name": "THANVI NISHANT JAGDISH", "Roll Number": "246" },
    { "Name": "TIWARI ALOK VINAY", "Roll Number": "247" },
    { "Name": "TOGAR STEAV KEPHA", "Roll Number": "248" },
    { "Name": "YADAV ADITYA ARVIND", "Roll Number": "249" },
    { "Name": "SINGH PRIYANSHU DEVENDRA", "Roll Number": "250" },
    { "Name": "TEMBHARE ASHWIN PREMLAL", "Roll Number": "251" }
]

# --- Session State ---
if "present_list" not in st.session_state:
    st.session_state.present_list = []

if "absent_list" not in st.session_state:
    st.session_state.absent_list = []

st.title("📋 Class Attendance System")
st.subheader("Mark Attendance Below")

# --- Attendance Form ---
with st.form("attendance_form"):
    marked_attendance = {}
    for student in student_list:
        key = f"{student['Roll Number']}_status"
        status = st.checkbox(
            f"{student['Roll Number']} - {student['Name']}",
            key=key,
            value=True
        )
        marked_attendance[student["Roll Number"]] = "Present" if status else "Absent"

    submitted = st.form_submit_button("✅ Submit Attendance")

    if submitted:
        today = datetime.now().strftime("%Y-%m-%d")
        st.session_state.present_list.clear()
        st.session_state.absent_list.clear()

        for student in student_list:
            record = {
                "Date": today,
                "Name": student["Name"],
                "Roll Number": student["Roll Number"]
            }
            if marked_attendance[student["Roll Number"]] == "Present":
                st.session_state.present_list.append(record)
            else:
                st.session_state.absent_list.append(record)

        st.success("Attendance submitted successfully!")

# --- Display Tables ---
def show_table(title, data_list):
    if data_list:
        st.subheader(title)
        df = pd.DataFrame(data_list)
        st.dataframe(df, use_container_width=True)
        return df
    return pd.DataFrame()

df_present = show_table("Present Students", st.session_state.present_list)
df_absent = show_table("Absent Students", st.session_state.absent_list)

# --- Create Word file ---
def create_word_file(present_df, absent_df):
    doc = Document()
    doc.add_heading("Class Attendance Report", 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph("")

    if not present_df.empty:
        doc.add_heading("Present Students", level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Date"
        hdr_cells[1].text = "Roll Number"
        hdr_cells[2].text = "Name"
        for _, row in present_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row["Date"]
            row_cells[1].text = str(row["Roll Number"])
            row_cells[2].text = row["Name"]

    if not absent_df.empty:
        doc.add_heading("Absent Students", level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Date"
        hdr_cells[1].text = "Roll Number"
        hdr_cells[2].text = "Name"
        for _, row in absent_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row["Date"]
            row_cells[1].text = str(row["Roll Number"])
            row_cells[2].text = row["Name"]

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Create PDF file ---
def create_pdf_file(present_df, absent_df):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", "B", 16)

    pdf.cell(200, 10, txt="Class Attendance Report", ln=True, align="C")
    pdf.set_font("Arial", "", 12)
    pdf.cell(200, 10, txt=f"Date: {datetime.now().strftime('%Y-%m-%d')}", ln=True, align="L")
    pdf.ln(10)

    if not present_df.empty:
        pdf.set_font("Arial", "B", 14)
        pdf.cell(200, 10, txt="Present Students", ln=True, align="L")
        pdf.set_font("Arial", "", 12)
        for _, row in present_df.iterrows():
            pdf.cell(200, 10, txt=f"{row['Roll Number']} - {row['Name']}", ln=True)

    if not absent_df.empty:
        pdf.set_font("Arial", "B", 14)
        pdf.cell(200, 10, txt="Absent Students", ln=True, align="L")
        pdf.set_font("Arial", "", 12)
        for _, row in absent_df.iterrows():
            pdf.cell(200, 10, txt=f"{row['Roll Number']} - {row['Name']}", ln=True)

    # ✅ Fix: Use string output + BytesIO
    pdf_output = pdf.output(dest='S').encode('latin1')
    return BytesIO(pdf_output)


# --- Download Buttons ---
if submitted:
    word_file = create_word_file(df_present, df_absent)
    pdf_file = create_pdf_file(df_present, df_absent)

    st.download_button(
        label="📥 Download Word File (.docx)",
        data=word_file,
        file_name="attendance_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.download_button(
        label="📄 Download PDF File (.pdf)",
        data=pdf_file,
        file_name="attendance_report.pdf",
        mime="application/pdf"
    )
